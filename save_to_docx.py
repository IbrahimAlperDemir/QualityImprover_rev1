from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from form_tracker import get_next_form_number  # 🔄 Yeni eklendi

def save_to_docx(text: str, filename: str = "Gereksinim_Dokumani.docx", doc_type: str = "PRD") -> None:
    doc = Document()

    # 📌 Otomatik Form Numarası ve Header
    form_code = get_next_form_number(doc_type=doc_type)  # Örn: FRM-PRD-004-A
    header_text = f"{form_code} | Ürün Gereksinim Dokümanı | Versiyon: 1.0 | Tarih: {datetime.today().strftime('%d.%m.%Y')}"

    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run(header_text)
    run.bold = True
    run.font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 📄 Başlık
    title = doc.add_heading("Ürün Özellik Gereksinim Dokümanı", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    # 📝 İçerik
    for line in text.splitlines():
        clean_line = line.strip()
        if not clean_line:
            doc.add_paragraph()
        elif clean_line.startswith(tuple("1234567890")) and clean_line[1:3] == ". ":
            para = doc.add_paragraph()
            run = para.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(12)
        elif clean_line.endswith(":"):
            para = doc.add_paragraph()
            run = para.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(11)
        else:
            para = doc.add_paragraph(clean_line)
            para.paragraph_format.space_after = Pt(6)

    # 💾 Kaydet
    doc.save(filename)
    print(f"✅ '{filename}' başarıyla kaydedildi.")
